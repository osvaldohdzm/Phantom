"""Pruebas del generador Word."""

import uuid
from datetime import datetime, timezone

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

from app.models.core import Finding, Severity
from app.services.docx_generator import (
    METODO_PLACEHOLDER,
    _apply_cyb001_table_formatting,
    _eliminar_ultimas_filas_si_es_salida_prueba_seguridad,
    _format_cell_content_bold,
    _format_cell_content_plain,
    _format_security_detail_content_cells,
    _format_security_detail_header_cells_white,
    _format_table_body_content_plain,
    _replace_in_paragraph,
    _set_paragraph_formatted_text,
    build_finding_replacements,
)
from app.services.report_text_preprocess import clean_tool_output


def _run_color_val(run):
    r_pr = run._element.find(qn("w:rPr"))
    if r_pr is None:
        return None
    color = r_pr.find(qn("w:color"))
    if color is None:
        return None
    return color.get(qn("w:val"))


def test_security_detail_content_uses_auto_black_not_bold():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "DETALLE DE PRUEBAS DE SEGURIDAD"
    table.rows[0].cells[1].text = "DETALLE DE PRUEBAS DE SEGURIDAD"
    table.rows[1].cells[0].text = "Escaneo automatizado"
    table.rows[1].cells[1].text = "Plugin output here"

    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    _format_security_detail_content_cells(doc)

    header_run = table.rows[0].cells[0].paragraphs[0].runs[0]
    assert header_run.font.color.rgb == RGBColor(0xFF, 0xFF, 0xFF)
    assert _run_is_bold(header_run)

    for content_run in (
        table.rows[1].cells[0].paragraphs[0].runs[0],
        table.rows[1].cells[1].paragraphs[0].runs[0],
    ):
        assert _run_color_val(content_run) == "000000"
        assert not _run_is_bold(content_run)


def test_format_cell_content_bold_nested_table():
    doc = Document()
    outer = doc.add_table(rows=1, cols=1)
    cell = outer.rows[0].cells[0]
    nested = cell.add_table(rows=1, cols=1)
    nested.rows[0].cells[0].text = "nested text"
    run = nested.rows[0].cells[0].paragraphs[0].runs[0]
    run.bold = False

    _format_cell_content_bold(cell)

    nested_run = nested.rows[0].cells[0].paragraphs[0].runs[0]
    assert _run_color_val(nested_run) == "auto"
    assert _run_is_bold(nested_run)


def test_format_cell_auto_black_nested_table():
    doc = Document()
    outer = doc.add_table(rows=1, cols=1)
    cell = outer.rows[0].cells[0]
    nested = cell.add_table(rows=1, cols=1)
    nested.rows[0].cells[0].text = "nested text"
    run = nested.rows[0].cells[0].paragraphs[0].runs[0]
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    _format_cell_content_plain(cell)

    assert _run_color_val(nested.rows[0].cells[0].paragraphs[0].runs[0]) == "000000"


def _make_finding(**kwargs) -> Finding:
    defaults = {
        "id": uuid.uuid4(),
        "titulo": "Test finding",
        "severidad": Severity.medium,
        "created_at": datetime.now(timezone.utc),
    }
    defaults.update(kwargs)
    return Finding(**defaults)


def test_metodo_deteccion_markdown_preserved():
    finding = _make_finding(
        metodo_deteccion="**Escaneo manual** con validación\n![cap](data:image/png;base64,abc)",
    )
    replacements = build_finding_replacements(finding)
    assert replacements["«Tipo de texto de método de detección»"] == "markdown"
    assert "**Escaneo manual**" in replacements[METODO_PLACEHOLDER]


def test_explicacion_tecnica_uses_own_field_not_descripcion():
    finding = _make_finding(
        descripcion="Descripción del catálogo",
        explicacion_tecnica="**Explicación propia** del analista",
    )
    replacements = build_finding_replacements(finding)
    assert replacements["«Tipo de texto de explicación técnica»"] == "markdown"
    assert "**Explicación propia**" in replacements["«Explicación técnica»"]
    assert "Descripción del catálogo" not in replacements["«Explicación técnica»"]


def test_explicacion_tecnica_falls_back_to_descripcion_when_empty():
    finding = _make_finding(
        descripcion="Solo descripción disponible",
        explicacion_tecnica="",
    )
    replacements = build_finding_replacements(finding)
    assert "Solo descripción disponible" in replacements["«Explicación técnica»"]


def test_evidencia_placeholder_not_mapped_to_salidas():
    finding = _make_finding(raw_tool_output="plugin output line")
    replacements = build_finding_replacements(finding)
    assert "«Evidencia»" not in replacements


def test_detail_header_stays_white_on_blue():
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    cell = table.rows[0].cells[1]
    cell.text = "DETALLE DE PRUEBAS DE SEGURIDAD"
    run = cell.paragraphs[0].runs[0]
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    _format_security_detail_content_cells(doc)
    _format_security_detail_header_cells_white(doc)

    header_run = cell.paragraphs[0].runs[0]
    assert header_run.font.color.rgb == RGBColor(0xFF, 0xFF, 0xFF)
    assert _run_is_bold(header_run)


def test_merged_detalle_header_stays_white_after_full_formatting():
    """Fila fusionada (una celda): el encabezado DETALLE no debe volverse negro al final."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    header = table.rows[0].cells[0]
    header.text = "DETALLE DE PRUEBAS DE SEGURIDAD"
    header.merge(table.rows[0].cells[1])
    content = table.rows[1].cells[0]
    content.text = "Método de detección"
    content.merge(table.rows[1].cells[1])

    _apply_cyb001_table_formatting(doc, {}, "Medium", None, None)

    run = header.paragraphs[0].runs[0]
    assert run.font.color.rgb == RGBColor(0xFF, 0xFF, 0xFF)
    assert _run_color_val(run) == "FFFFFF"
    assert _run_is_bold(run)


def test_formatted_text_only_bolds_markdown_segments():
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("plantilla")
    run.bold = True
    _set_paragraph_formatted_text(
        para._element, "Texto normal y **solo esto** en negrita."
    )
    texts = [r.text for r in para.runs if r.text]
    joined = "".join(texts)
    assert "Texto normal" in joined
    bold_runs = [r for r in para.runs if r.text and r.bold]
    plain_runs = [r for r in para.runs if r.text and not r.bold]
    assert any("solo esto" in r.text for r in bold_runs)
    assert any("Texto normal" in r.text for r in plain_runs)


def test_detalle_header_is_not_mapped_to_salidas():
    finding = _make_finding(
        raw_tool_output="plugin output only",
        metodo_deteccion="Escaneo Nessus",
    )
    replacements = build_finding_replacements(finding)
    assert "«DETALLE DE PRUEBAS DE SEGURIDAD»" not in replacements
    assert replacements["«Salidas de herramienta»"] == replacements["«Salida de herramienta»"]
    assert "plugin output only" in replacements["«Salidas de herramienta»"]


def test_cyb009_cleans_salidas_in_replacements():
    finding = _make_finding(
        raw_tool_output="Nessus found issue at https://host/path/http://other",
    )
    replacements = build_finding_replacements(finding)
    salida = replacements["«Salidas de herramienta»"]
    assert "The scanner tool" in salida
    assert "Nessus" not in salida
    assert "/\nhttp://" in salida


def test_clean_tool_output_splits_slash_http():
    raw = "see https://example.com/path/http://other.example"
    cleaned = clean_tool_output(raw)
    assert "/\nhttp://other" in cleaned


def test_eliminar_filas_when_salida_and_metodo_empty():
    doc = Document()
    table = doc.add_table(rows=4, cols=2)
    for i in range(4):
        table.rows[i].cells[0].text = f"Label {i}"
        table.rows[i].cells[1].text = f"Value {i}"
    replacements = {
        "«Salidas de herramienta»": "",
        METODO_PLACEHOLDER: "",
    }
    _eliminar_ultimas_filas_si_es_salida_prueba_seguridad(doc, replacements)
    assert len(doc.tables[0].rows) == 2


def _run_is_bold(run) -> bool:
    if run.bold:
        return True
    r_pr = run._element.find(qn("w:rPr"))
    if r_pr is None:
        return False
    bold_el = r_pr.find(qn("w:b"))
    if bold_el is None:
        return False
    val = bold_el.get(qn("w:val"))
    return val is None or val.lower() not in ("0", "false")


def test_detail_content_cells_plain_after_replacement():
    doc = Document()
    table = doc.add_table(rows=4, cols=2)
    table.rows[0].cells[0].text = "AMENAZA"
    para = table.rows[0].cells[1].paragraphs[0]
    run = para.add_run("«Amenaza»")
    run.bold = True
    table.rows[1].cells[0].text = "DETALLE DE PRUEBAS DE SEGURIDAD"
    det = table.rows[1].cells[1].paragraphs[0]
    det_run = det.add_run("«Método de detección»")
    det_run.bold = True
    det.add_run("\n«Salidas de herramienta»")
    det.add_run("\n«Explicación técnica»")

    _replace_in_paragraph(
        table.rows[0].cells[1].paragraphs[0],
        {"«Amenaza»": "Un atacante podría explotar el sistema."},
        keep_bold=False,
    )
    _replace_in_paragraph(
        det,
        {
            "«Método de detección»": "Escaneo Nessus",
            "«Salidas de herramienta»": "plugin output",
            "«Explicación técnica»": "Detalle técnico",
        },
        keep_bold=False,
    )
    _format_security_detail_content_cells(doc)
    _format_table_body_content_plain(doc)

    amenaza_runs = table.rows[0].cells[1].paragraphs[0].runs
    assert amenaza_runs
    assert not any(_run_is_bold(r) for r in amenaza_runs if r.text.strip())

    detail_runs = table.rows[1].cells[1].paragraphs[0].runs
    assert detail_runs
    assert not any(_run_is_bold(r) for r in detail_runs if r.text.strip())


def test_detalle_outer_paragraphs_black_after_simulated_label_white():
    """Método/explicación (w:p exteriores) deben volver a negro aunque se pinten de blanco."""
    from app.services.docx_generator import (
        _format_label_cell_white,
        _format_security_detail_content_row_final,
    )

    doc = Document()
    table = doc.add_table(rows=2, cols=1)
    table.rows[0].cells[0].text = "DETALLE DE PRUEBAS DE SEGURIDAD"
    cell = table.rows[1].cells[0]
    cell.paragraphs[0].text = "Escaneo Nessus"
    nested = cell.add_table(rows=1, cols=1)
    nested.rows[0].cells[0].text = "HTTP/1.1 200 OK"
    cell.paragraphs[1].text = "Explicación con amenaza en sistemas afectados."

    _format_label_cell_white(cell)
    _format_security_detail_content_row_final(doc)

    for para in cell.paragraphs:
        for run in para.runs:
            if not run.text.strip():
                continue
            assert _run_color_val(run) == "000000"
            assert not _run_is_bold(run)
    nested_run = nested.rows[0].cells[0].paragraphs[0].runs[0]
    assert _run_color_val(nested_run) == "000000"


def test_detalle_content_not_whited_when_text_contains_label_keywords():
    """La última fila no debe volverse blanca si el texto menciona «afectados»/«amenaza»."""
    doc = Document()
    table = doc.add_table(rows=2, cols=1)
    table.rows[0].cells[0].text = "DETALLE DE PRUEBAS DE SEGURIDAD"
    content = table.rows[1].cells[0]
    content.text = (
        "Escaneo Nessus\n"
        "Los sistemas afectados muestran una amenaza crítica.\n"
        "Remediación pendiente."
    )
    replacements = {
        "«Salidas de herramienta»": "HTTP/1.1 200 OK",
        METODO_PLACEHOLDER: "Escaneo",
    }
    _apply_cyb001_table_formatting(doc, replacements, "High", None, None)
    for para in content.paragraphs:
        for run in para.runs:
            if not run.text.strip():
                continue
            assert _run_color_val(run) == "000000"
            assert not _run_is_bold(run)


def test_salidas_placeholder_in_nested_table_is_replaced():
    """Plantilla CYB001: «Salidas de herramienta» vive en tabla anidada bajo DETALLE."""
    from pathlib import Path

    from app.services.docx_generator import generate_single_finding_docx

    tpl = Path(__file__).resolve().parents[1] / (
        "storage/templates/789cb2c9-57fe-4d29-9d6e-715e2650f31d.docx"
    )
    if not tpl.is_file():
        return

    finding = Finding(
        id=uuid.uuid4(),
        titulo="Weak TLS",
        severidad=Severity.high,
        metodo_deteccion="Escaneo Nessus",
        explicacion_tecnica="Protocolo obsoleto",
        raw_tool_output="TLS 1.1 enabled on port 443",
        componente_afectado="10.0.0.1",
        created_at=datetime.now(timezone.utc),
    )
    replacements = build_finding_replacements(
        finding,
        merged_raw_output=(
            "-----[Salida correspondiente a: 10.0.0.1 ]-----\nTLS 1.1 enabled on port 443"
        ),
    )
    out = tpl.parent / "_pytest_nested_salidas.docx"
    try:
        generate_single_finding_docx(str(tpl), str(out), replacements, severity="High")
        doc = Document(str(out))
        detail_cell = doc.tables[0].rows[7].cells[0]
        nested_cell = detail_cell.tables[0].rows[0].cells[0]
        nested_text = nested_cell.text
        assert "«Salidas de herramienta»" not in nested_text
        assert "TLS 1.1 enabled on port 443" in nested_text
        assert "-----[Salida correspondiente a: 10.0.0.1 ]-----" in nested_text
        for para in nested_cell.paragraphs:
            for run in para.runs:
                if run.text.strip():
                    assert not _run_is_bold(run)
                    assert _run_color_val(run) == "000000"
        header_run = doc.tables[0].rows[6].cells[0].paragraphs[0].runs[0]
        assert header_run.bold
        assert _run_color_val(header_run) == "FFFFFF"
    finally:
        if out.is_file():
            out.unlink()


def test_eliminar_nested_table_when_salida_empty_metodo_present():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Método de detección"
    table.rows[0].cells[1].text = "Escaneo"
    cell = table.rows[1].cells[1]
    cell.text = ""
    nested = cell.add_table(rows=1, cols=1)
    nested.rows[0].cells[0].text = "gray box salida"
    replacements = {
        "«Salidas de herramienta»": "",
        METODO_PLACEHOLDER: "Escaneo automatizado",
    }
    _eliminar_ultimas_filas_si_es_salida_prueba_seguridad(doc, replacements)
    assert len(table.rows[1].cells[1].tables) == 0
