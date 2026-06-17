"""Genera Tabla de hallazgos (.docx) — resumen de severidad, nombre y componentes."""

from __future__ import annotations

import tempfile
import uuid
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.models.core import Asset, Finding
from app.services.docx_generator import _format_risk_cell, _set_cell_shading
from app.services.finding_grouping import (
    first_non_empty_field,
    prepare_grouped_rows_for_report,
)
from app.services.report_text_preprocess import preprocess_report_field

HEADER_FILL = "4472C4"
HEADER_FONT = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_FONT_NAME = "Roboto"
TABLE_FONT_SIZE = Pt(12)
COLUMN_HEADERS = ("Severidad", "Nombre de vulnerabilidad", "Componente afectado")


def _set_run_typography(
    run, *, bold: Optional[bool] = None, color: Optional[RGBColor] = None
) -> None:
    run.font.name = TABLE_FONT_NAME
    run.font.size = TABLE_FONT_SIZE
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), TABLE_FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), TABLE_FONT_NAME)
    r_fonts.set(qn("w:cs"), TABLE_FONT_NAME)


def _format_header_cell(cell, text: str) -> None:
    _set_cell_shading(cell, HEADER_FILL)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in list(para.runs):
        run.text = ""
    run = para.add_run(text)
    _set_run_typography(run, bold=True, color=HEADER_FONT)


def _set_paragraph_text(cell, text: str, *, align: WD_ALIGN_PARAGRAPH) -> None:
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    para.alignment = align
    for run in list(para.runs):
        run.text = ""
    if text:
        run = para.add_run(text)
        _set_run_typography(run, bold=False)


def _format_severity_cell(cell, severity: str) -> None:
    """Negrita + relleno INAI por nivel (vía docx_generator) y tipografía Roboto 12."""
    _format_risk_cell(cell, severity)
    for para in cell.paragraphs:
        for run in para.runs:
            _set_run_typography(run, bold=True)


def _format_vulnerability_name(finding: Finding, group_members: list[Finding]) -> str:
    titulo = preprocess_report_field(finding.titulo or "")
    cve = (first_non_empty_field(group_members, "cve") or finding.cve or "").strip()
    if cve and cve.upper() not in titulo.upper():
        label = cve if cve.upper().startswith("CVE") else f"CVE-{cve}"
        titulo = f"{titulo} ({label})"
    return titulo


def _apply_column_widths(table) -> None:
    widths = (Inches(1.15), Inches(3.85), Inches(2.5))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width


def generate_findings_table_docx(
    findings: list[Finding],
    assets_map: Optional[dict[uuid.UUID, Asset]] = None,
    *,
    output_path: Optional[str] = None,
) -> str:
    """Crea un .docx con tabla de hallazgos agrupados. Devuelve la ruta del archivo."""
    grouped = prepare_grouped_rows_for_report(findings, assets_map)

    doc = Document()
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_row = table.rows[0]
    for cell, label in zip(header_row.cells, COLUMN_HEADERS):
        _format_header_cell(cell, label)

    for row_data in grouped:
        finding = row_data.representative
        row = table.add_row()

        sev_cell = row.cells[0]
        sev_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        sev = finding.severidad.value if finding.severidad else "Medium"
        _format_severity_cell(sev_cell, sev)
        for para in sev_cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        name_cell = row.cells[1]
        name_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_paragraph_text(
            name_cell,
            _format_vulnerability_name(finding, row_data.members),
            align=WD_ALIGN_PARAGRAPH.LEFT,
        )

        comp_cell = row.cells[2]
        comp_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_paragraph_text(
            comp_cell,
            row_data.merged_componente_afectado,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    _apply_column_widths(table)

    if output_path:
        path = Path(output_path)
    else:
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", prefix="tabla_hallazgos_", delete=False)
        path = Path(tmp.name)
        tmp.close()

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return str(path)
