"""Generador de reportes Word desde plantillas con marcadores «Columna».

Replica la lógica del macro VBA CYB001_GenerarDocumentosVulnerabilidadesWord:
- Sustitución de marcadores «Nombre de columna» por valores del hallazgo
- Un documento por hallazgo + fusión en consolidado
- Soporte básico de markdown en «Explicación técnica»
"""

from __future__ import annotations

import base64
import re
import shutil
import tempfile
import uuid
import zipfile
from copy import deepcopy
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any, Literal, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor
from docx.text.paragraph import Paragraph

from app.models.core import Finding, Asset, Engagement, VulnerabilityCatalog
from app.services.finding_grouping import (
    first_non_empty_field,
    prepare_grouped_rows_for_report,
)
from app.services.vulns_catalog_lookup import catalog_text
from app.services.report_text_preprocess import (
    cyb009_salidas_herramienta,
    decode_html_entities,
    normalize_line_breaks,
    normalize_severity_label,
    preprocess_report_field,
    split_plain_lines,
    strip_html_tags,
)

PLACEHOLDER_RE = re.compile(r"«([^»]+)»")
RISK_COLORS = {
    "critical": RGBColor(0xEF, 0x44, 0x44),
    "high": RGBColor(0xF9, 0x73, 0x16),
    "medium": RGBColor(0xEA, 0xB3, 0x08),
    "low": RGBColor(0x3B, 0x82, 0xF6),
    "info": RGBColor(0x6B, 0x72, 0x80),
}
FONT_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
FONT_BLACK = RGBColor(0x00, 0x00, 0x00)
LABEL_WHITE_KEYWORDS = (
    "DESCRIPCI",
    "AMENAZA",
    "REMEDIACI",
    "REFERENCI",
    "SISTEMA",
    "RUTA",
    "AFECTAD",
    "SEVERIDAD",
    "DETALLE DE PRUEBAS",
)
SEVERITY_ES = {
    "Critical": "Crítico",
    "High": "Alto",
    "Medium": "Medio",
    "Low": "Bajo",
    "Info": "Informativo",
}

# Colores INAI (fondo hex + color de fuente) — FormatearCeldaNivelRiesgo
SEVERITY_INAI_STYLE: dict[str, tuple[str, RGBColor]] = {
    "CRÍTICA": ("7030A0", FONT_WHITE),
    "ALTA": ("FF0000", FONT_WHITE),
    "MEDIA": ("FFFF00", FONT_BLACK),
    "BAJA": ("00B050", FONT_WHITE),
    "INFORMATIVA": ("E7E6E6", FONT_BLACK),
}

def extract_placeholders_from_docx(template_path: str) -> list[str]:
    """Extrae nombres de marcadores «...» del documento."""
    found: set[str] = set()
    doc = Document(template_path)
    for para in _iter_all_paragraphs(doc):
        for match in PLACEHOLDER_RE.finditer(para.text):
            found.add(match.group(1))
    return sorted(found)


def _iter_cell_paragraphs(cell):
    """Párrafos de una celda, incluyendo tablas anidadas (p. ej. «Salidas de herramienta»)."""
    for para in cell.paragraphs:
        yield para
    for nested in cell.tables:
        for row in nested.rows:
            for nested_cell in row.cells:
                yield from _iter_cell_paragraphs(nested_cell)


def _iter_all_paragraphs(doc: Document):
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_cell_paragraphs(cell)


def _replace_in_paragraph(
    paragraph, replacements: dict[str, str], *, keep_bold: bool = True
) -> None:
    full_text = paragraph.text
    if not full_text or "«" not in full_text:
        return
    new_text = full_text
    for key, value in replacements.items():
        new_text = new_text.replace(key, value or "")
    if new_text == full_text:
        return
    _set_paragraph_text_keep_markup(
        paragraph._element, new_text, keep_bold=keep_bold
    )


def _replace_placeholders_in_doc(
    doc: Document, replacements: dict[str, str], *, keep_bold: bool = True
) -> None:
    for para in _iter_all_paragraphs(doc):
        _replace_in_paragraph(para, replacements, keep_bold=keep_bold)


def _is_markdown(text: Optional[str]) -> bool:
    if not text:
        return False
    markers = ["##", "**", "```", "- ", "* ", "[", "](", "# ", "![", "data:image"]
    return any(m in text for m in markers)


IMAGE_MD_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
BOLD_SEGMENT_RE = re.compile(r"\*\*([^*]+)\*\*")

EXPLICACION_PLACEHOLDER = "«Explicación técnica»"
METODO_PLACEHOLDER = "«Método de detección»"
SALIDAS_PLACEHOLDER_KEYS = (
    "«Salidas de herramienta»",
    "«Salida de herramienta»",
)
DETAIL_CONTENT_PLACEHOLDER_KEYS = (
    METODO_PLACEHOLDER,
    EXPLICACION_PLACEHOLDER,
    *SALIDAS_PLACEHOLDER_KEYS,
)
METADATA_PLACEHOLDER_PREFIX = "«Tipo de texto"


@dataclass
class MarkdownBlock:
    kind: Literal["heading1", "heading2", "bullet", "paragraph", "image"]
    text: str = ""
    alt: str = ""
    src: str = ""


def _parse_markdown_blocks(markdown_text: str) -> list[MarkdownBlock]:
    blocks: list[MarkdownBlock] = []
    for raw_line in markdown_text.split("\n"):
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            continue

        img_match = IMAGE_MD_RE.search(stripped)
        if img_match and stripped.strip().startswith("!["):
            blocks.append(
                MarkdownBlock(
                    kind="image",
                    alt=img_match.group(1),
                    src=img_match.group(2),
                )
            )
            continue

        if stripped.startswith("## "):
            blocks.append(MarkdownBlock(kind="heading2", text=stripped[3:]))
        elif stripped.startswith("# "):
            blocks.append(MarkdownBlock(kind="heading1", text=stripped[2:]))
        elif stripped.startswith("- ") or stripped.startswith("* "):
            blocks.append(MarkdownBlock(kind="bullet", text=stripped[2:]))
        else:
            blocks.append(MarkdownBlock(kind="paragraph", text=line))

    return blocks


def _load_image_stream(src: str) -> Optional[BytesIO]:
    src = src.strip()
    if src.startswith("data:"):
        try:
            _header, payload = src.split(",", 1)
            return BytesIO(base64.b64decode(payload))
        except (ValueError, base64.binascii.Error):
            return None
    path = Path(src)
    if path.is_file():
        return BytesIO(path.read_bytes())
    return None


def _markdown_block_text(block: MarkdownBlock) -> str:
    """Texto del bloque conservando **negritas** para renderizado en Word."""
    return block.text.strip()


def _preserve_markdown_field(text: Optional[str]) -> str:
    """Limpieza mínima que conserva markdown (negritas, imágenes)."""
    if not text:
        return ""
    t = strip_html_tags(str(text))
    t = normalize_line_breaks(t)
    return decode_html_entities(t)


def _insert_image_block(container, block: MarkdownBlock, template_para: Paragraph) -> Paragraph:
    """Inserta imagen centrada; el pie de foto hereda el estilo del párrafo plantilla."""
    para = container.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stream = _load_image_stream(block.src)
    if stream:
        try:
            run = para.add_run()
            run.add_picture(stream, width=Inches(5.2))
        except Exception:
            para.add_run(f"[Imagen: {block.alt or 'evidencia'}]")
    else:
        para.add_run(f"[Imagen no disponible: {block.alt or block.src[:80]}]")
    if block.alt:
        cap = deepcopy(template_para._element)
        _set_paragraph_text_keep_markup(cap, block.alt)
        para._element.addnext(cap)
    return para


def _strip_bold_from_r_pr(r_pr) -> None:
    if r_pr is None:
        return
    bold_el = r_pr.find(qn("w:b"))
    if bold_el is not None:
        r_pr.remove(bold_el)


def _clone_r_pr_plain(template_r_pr) -> Optional[Any]:
    if template_r_pr is None:
        return None
    cloned = deepcopy(template_r_pr)
    _strip_bold_from_r_pr(cloned)
    return cloned


def _set_paragraph_text_keep_markup(p_elem, text: str, *, keep_bold: bool = True) -> None:
    """Sustituye solo el texto del párrafo conservando pPr/rPr de la plantilla."""
    value = text.strip()
    runs = p_elem.findall(qn("w:r"))
    if not runs:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = value
        r.append(t)
        p_elem.append(r)
        return
    for idx, run in enumerate(runs):
        if not keep_bold:
            _strip_bold_from_r_pr(run.find(qn("w:rPr")))
        t_nodes = run.findall(qn("w:t"))
        if idx == 0:
            if t_nodes:
                t_nodes[0].text = value
                for extra in t_nodes[1:]:
                    extra.text = ""
            else:
                t = OxmlElement("w:t")
                t.text = value
                run.append(t)
        else:
            for t in t_nodes:
                t.text = ""
            for child in list(run):
                if child.tag != qn("w:t"):
                    run.remove(child)


def _set_paragraph_formatted_text(p_elem, text: str) -> None:
    """Sustituye texto del párrafo renderizando segmentos **negrita**."""
    if "**" not in text:
        _set_paragraph_text_keep_markup(p_elem, text, keep_bold=False)
        return

    runs = p_elem.findall(qn("w:r"))
    template_r_pr = None
    if runs:
        r_pr = runs[0].find(qn("w:rPr"))
        if r_pr is not None:
            template_r_pr = r_pr

    for run in list(runs):
        p_elem.remove(run)

    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        content = part[2:-2] if is_bold else part

        run = OxmlElement("w:r")
        if is_bold:
            if template_r_pr is not None:
                run.append(deepcopy(template_r_pr))
            r_pr = run.find(qn("w:rPr"))
            if r_pr is None:
                r_pr = OxmlElement("w:rPr")
                run.insert(0, r_pr)
            bold_el = r_pr.find(qn("w:b"))
            if bold_el is None:
                bold_el = OxmlElement("w:b")
                r_pr.append(bold_el)
            bold_el.set(qn("w:val"), "true")
        else:
            plain_r_pr = _clone_r_pr_plain(template_r_pr)
            if plain_r_pr is not None:
                run.append(plain_r_pr)
        text_el = OxmlElement("w:t")
        text_el.set(qn("xml:space"), "preserve")
        text_el.text = content
        run.append(text_el)
        p_elem.append(run)


def _find_placeholder_in_table(
    table, placeholder_key: str
) -> tuple[Any, Paragraph] | tuple[None, None]:
    """Busca marcador en tabla (recursivo en tablas anidadas)."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if placeholder_key in para.text:
                    return cell, para
            for nested in cell.tables:
                found = _find_placeholder_in_table(nested, placeholder_key)
                if found[0] is not None:
                    return found
    return None, None


def _find_placeholder_in_doc(
    doc: Document, placeholder_key: str
) -> tuple[Any, Paragraph] | tuple[None, None]:
    """Busca celda/párrafo con un marcador «…»."""
    for table in doc.tables:
        found = _find_placeholder_in_table(table, placeholder_key)
        if found[0] is not None:
            return found
    for para in doc.paragraphs:
        if placeholder_key in para.text:
            return doc, para
    return None, None


def _insert_rich_markdown_at_placeholder(
    doc: Document, placeholder_key: str, markdown_text: str
) -> None:
    """Reemplaza un marcador por bloques markdown (negritas, imágenes centradas)."""
    container, anchor = _find_placeholder_in_doc(doc, placeholder_key)
    if container is None or anchor is None:
        return

    blocks = _parse_markdown_blocks(markdown_text)
    if not blocks:
        _replace_in_paragraph(anchor, {placeholder_key: ""})
        return

    anchor_text = anchor.text.replace(placeholder_key, "").strip()
    _set_paragraph_text_keep_markup(anchor._element, anchor_text)

    anchor_elem = anchor._element
    anchor_free = not anchor_text
    last_elem = anchor_elem

    for block in blocks:
        if block.kind == "image":
            img_para = _insert_image_block(container, block, anchor)
            last_elem = img_para._element
            anchor_free = False
            continue

        formatted = _markdown_block_text(block)
        if not formatted:
            continue

        if anchor_free:
            _set_paragraph_formatted_text(anchor_elem, formatted)
            anchor_free = False
            last_elem = anchor_elem
            continue

        new_p = deepcopy(anchor_elem)
        _set_paragraph_formatted_text(new_p, formatted)
        last_elem.addnext(new_p)
        last_elem = new_p


def _insert_rich_markdown_explicacion(doc: Document, markdown_text: str) -> None:
    """Reemplaza «Explicación técnica» conservando estilo de la plantilla."""
    _insert_rich_markdown_at_placeholder(doc, EXPLICACION_PLACEHOLDER, markdown_text)


def _insert_rich_markdown_metodo(doc: Document, markdown_text: str) -> None:
    """Reemplaza «Método de detección» conservando estilo de la plantilla."""
    _insert_rich_markdown_at_placeholder(doc, METODO_PLACEHOLDER, markdown_text)


def _insert_rich_markdown_salidas(doc: Document, markdown_text: str) -> None:
    """Inserta markdown en marcadores de salida de herramienta / detalle de pruebas."""
    for key in SALIDAS_PLACEHOLDER_KEYS:
        _, anchor = _find_placeholder_in_doc(doc, key)
        if anchor is not None:
            _insert_rich_markdown_at_placeholder(doc, key, markdown_text)


def _clear_paragraph_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        run.text = ""


def _get_header_cells(doc: Document) -> tuple[Any, Any]:
    """Fila 0: primera celda fusionada = nombre, siguiente distinta = severidad."""
    if not doc.tables or not doc.tables[0].rows:
        return None, None
    row0 = doc.tables[0].rows[0]
    if not row0.cells:
        return None, None
    title_cell = row0.cells[0]
    title_tc = id(title_cell._tc)
    severity_cell = None
    for cell in row0.cells[1:]:
        if id(cell._tc) != title_tc:
            severity_cell = cell
            break
    return title_cell, severity_cell


def _cell_is_label_header(cell) -> bool:
    """Etiqueta de fila (columna izquierda), no contenido narrativo bajo DETALLE."""
    text = cell.text.strip().upper()
    if not text or "\n" in text or len(text) > 120:
        return False
    return any(kw in text for kw in LABEL_WHITE_KEYWORDS)


def _is_detalle_header_cell(cell) -> bool:
    return "DETALLE DE PRUEBAS" in _cell_label_text(cell)


def _force_run_white(run) -> None:
    """Texto blanco forzado en XML (evita «auto»/tema que se ve negro sobre azul)."""
    run.bold = True
    run.font.color.rgb = FONT_WHITE
    r_pr = run._element.get_or_add_rPr()
    for child in list(r_pr):
        if child.tag in (qn("w:color"), qn("w:highlight")):
            r_pr.remove(child)
    bold_el = r_pr.find(qn("w:b"))
    if bold_el is None:
        bold_el = OxmlElement("w:b")
        r_pr.append(bold_el)
    bold_el.set(qn("w:val"), "true")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "FFFFFF")
    r_pr.append(color)


def _format_label_cell_white(cell) -> None:
    if cell is None:
        return
    for para in cell.paragraphs:
        for run in para.runs:
            _force_run_white(run)


def _set_run_font_auto(run) -> None:
    """Color automático de Word (negro por defecto sobre fondo claro)."""
    r_pr = run._element.get_or_add_rPr()
    for child in list(r_pr):
        if child.tag == qn("w:color"):
            r_pr.remove(child)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "auto")
    r_pr.append(color)


def _set_paragraph_default_run_plain_black(p_elem) -> None:
    """Quita negrita/blanco del rPr por defecto del párrafo (pPr/rPr)."""
    p_pr = p_elem.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        p_elem.insert(0, p_pr)
    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is not None:
        p_pr.remove(p_style)
    r_pr = p_pr.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        p_pr.append(r_pr)
    for tag in (qn("w:b"), qn("w:bCs")):
        bold_el = r_pr.find(tag)
        if bold_el is not None:
            r_pr.remove(bold_el)
    for child in list(r_pr):
        if child.tag == qn("w:color"):
            r_pr.remove(child)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    r_pr.append(color)


def _force_run_element_plain_black(r_elem) -> None:
    """Negrita/color en XML de un w:r (párrafos exteriores de la celda DETALLE)."""
    r_pr = r_elem.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        r_elem.insert(0, r_pr)
    for tag in (qn("w:b"), qn("w:bCs")):
        bold_el = r_pr.find(tag)
        if bold_el is not None:
            r_pr.remove(bold_el)
    for child in list(r_pr):
        if child.tag in (qn("w:color"), qn("w:highlight")):
            r_pr.remove(child)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    r_pr.append(color)


def _format_p_element_plain_black(p_elem) -> None:
    """Párrafo w:p directo en la celda: negro explícito, sin estilo heredado."""
    _set_paragraph_default_run_plain_black(p_elem)
    for r_elem in p_elem.findall(qn("w:r")):
        _force_run_element_plain_black(r_elem)


def _force_run_plain_black(run) -> None:
    """Texto negro explícito, sin negrita (última fila DETALLE / salidas Nessus)."""
    run.bold = False
    run.font.color.rgb = FONT_BLACK
    r_pr = run._element.get_or_add_rPr()
    bold_el = r_pr.find(qn("w:b"))
    if bold_el is not None:
        r_pr.remove(bold_el)
    for child in list(r_pr):
        if child.tag in (qn("w:color"), qn("w:highlight")):
            r_pr.remove(child)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    r_pr.append(color)


def _strip_bold_from_cell(cell) -> None:
    """Quita negrita de todo el texto de contenido (no etiquetas de fila)."""
    if cell is None:
        return
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = False
            _strip_bold_from_r_pr(run._element.find(qn("w:rPr")))
    for nested in cell.tables:
        for row in nested.rows:
            for nested_cell in row.cells:
                _strip_bold_from_cell(nested_cell)


def _force_run_bold(run) -> None:
    """Negrita forzada en XML (contenido DETALLE: método, salidas, explicación)."""
    run.bold = True
    r_pr = run._element.get_or_add_rPr()
    bold_el = r_pr.find(qn("w:b"))
    if bold_el is None:
        bold_el = OxmlElement("w:b")
        r_pr.append(bold_el)
    bold_el.set(qn("w:val"), "true")


def _format_cell_content_bold(cell) -> None:
    """Contenido DETALLE sustituido: color automático y negrita."""
    if cell is None:
        return
    for para in cell.paragraphs:
        for run in para.runs:
            _force_run_bold(run)
            _set_run_font_auto(run)
    for nested in cell.tables:
        for row in nested.rows:
            for nested_cell in row.cells:
                _format_cell_content_bold(nested_cell)


def _format_cell_content_plain(cell) -> None:
    """Contenido sustituido: negro explícito y sin negrita heredada de etiquetas."""
    if cell is None:
        return
    for child in cell._tc:
        if child.tag == qn("w:p"):
            _format_p_element_plain_black(child)
    for nested in cell.tables:
        for row in nested.rows:
            for nested_cell in row.cells:
                _format_cell_content_plain(nested_cell)


def _format_cell_auto_black(cell) -> None:
    """Alias retrocompatible: contenido plano en celdas de valor."""
    _format_cell_content_plain(cell)


def _format_security_detail_content_cells(doc: Document) -> None:
    """Contenido bajo «DETALLE DE PRUEBAS DE SEGURIDAD»: texto normal, color automático."""
    if not doc.tables:
        return
    table = doc.tables[0]
    detail_row_idx = _find_detalle_header_row_idx(table)
    if detail_row_idx is None:
        return

    for row_idx, row in enumerate(table.rows):
        if detail_row_idx is None or row_idx < detail_row_idx or not row.cells:
            continue
        if row_idx == detail_row_idx:
            for cell in row.cells:
                if _is_detalle_header_cell(cell):
                    _format_label_cell_white(cell)
            continue
        # Fila de contenido bajo el encabezado de sección: sin negrita, color auto.
        for cell in row.cells:
            if _is_detalle_header_cell(cell):
                _format_label_cell_white(cell)
            else:
                _format_cell_content_plain(cell)


def _format_table_body_content_plain(doc: Document) -> None:
    """Columna derecha: texto normal; bloque DETALLE (valores): sin negrita, color auto."""
    if not doc.tables:
        return
    table = doc.tables[0]
    title_cell, severity_cell = _get_header_cells(doc)
    title_tc = id(title_cell._tc) if title_cell is not None else None
    severity_tc = id(severity_cell._tc) if severity_cell is not None else None
    detail_row_idx = _find_detalle_header_row_idx(table)

    for row_idx, row in enumerate(table.rows):
        if not row.cells:
            continue
        is_detail_header_row = detail_row_idx is not None and row_idx == detail_row_idx
        is_detail_content_row = detail_row_idx is not None and row_idx > detail_row_idx

        if is_detail_header_row:
            continue

        if is_detail_content_row:
            for cell in row.cells:
                if _is_detalle_header_cell(cell):
                    continue
                _format_cell_content_plain(cell)
            continue

        for cell in row.cells[1:]:
            if severity_tc is not None and id(cell._tc) == severity_tc:
                continue
            if title_tc is not None and id(cell._tc) == title_tc:
                continue
            _format_cell_content_plain(cell)


def _row_is_label_value_pair(row) -> bool:
    """Fila etiqueta|valor (dos columnas distintas). No filas fusionadas de contenido."""
    if len(row.cells) < 2:
        return False
    return id(row.cells[0]._tc) != id(row.cells[1]._tc)


def _format_table_label_cells(table, *, through_row_idx: Optional[int] = None) -> list:
    """Texto blanco y negrita solo en columna izquierda de filas etiqueta|valor."""
    white_cells: list = []
    detalle_idx = _find_detalle_header_row_idx(table)
    stop_idx = through_row_idx
    if detalle_idx is not None:
        stop_idx = detalle_idx if stop_idx is None else min(stop_idx, detalle_idx)
    for row_idx, row in enumerate(table.rows):
        if stop_idx is not None and row_idx > stop_idx:
            break
        if not row.cells or not _row_is_label_value_pair(row):
            continue
        left = row.cells[0]
        if _cell_is_label_header(left):
            _format_label_cell_white(left)
            white_cells.append(left)
    return white_cells


def _format_security_detail_header_cells_white(doc: Document) -> None:
    """Encabezado azul «DETALLE DE PRUEBAS DE SEGURIDAD»: texto blanco y negrita."""
    if not doc.tables:
        return
    for row in doc.tables[0].rows:
        for cell in row.cells:
            if _is_detalle_header_cell(cell):
                _format_label_cell_white(cell)


def _find_detalle_header_row_idx(table) -> Optional[int]:
    for idx, row in enumerate(table.rows):
        if row.cells and "DETALLE DE PRUEBAS" in _cell_label_text(row.cells[0]):
            return idx
    return None


def _format_security_detail_content_row_final(doc: Document) -> None:
    """Último paso: fila bajo DETALLE siempre negro y sin negrita (incl. tablas anidadas)."""
    if not doc.tables:
        return
    table = doc.tables[0]
    detail_row_idx = _find_detalle_header_row_idx(table)
    if detail_row_idx is None:
        return
    for row in table.rows[detail_row_idx + 1 :]:
        seen_tc: set[int] = set()
        for cell in row.cells:
            tc_id = id(cell._tc)
            if tc_id in seen_tc:
                continue
            seen_tc.add(tc_id)
            if _is_detalle_header_cell(cell):
                continue
            # Párrafos exteriores (método / explicación) viven en w:p del tc, no en la subtabla.
            for child in cell._tc:
                if child.tag == qn("w:p"):
                    _format_p_element_plain_black(child)
            for nested in cell.tables:
                for nrow in nested.rows:
                    for nested_cell in nrow.cells:
                        _format_cell_content_plain(nested_cell)


def _replace_cell_text_preserve_format(cell, text: str) -> None:
    """Sustituye texto conservando estilo de la plantilla (tamaño, fuente, negrita)."""
    if cell is None or not cell.paragraphs:
        return
    _set_paragraph_text_keep_markup(cell.paragraphs[0]._element, text or "", keep_bold=True)
    for extra in cell.paragraphs[1:]:
        _clear_paragraph_runs(extra)


def _expand_cell_lines_from_template(cell, *, split_inline_bullets: bool = False) -> None:
    """Un ítem por línea clonando el párrafo plantilla (estilo, tamaño, viñetas)."""
    if not cell.paragraphs:
        return
    template_para = cell.paragraphs[0]
    lines = split_plain_lines(cell.text, split_inline_bullets=split_inline_bullets)
    if not lines:
        return

    _set_paragraph_text_keep_markup(template_para._element, lines[0], keep_bold=False)

    while len(cell.paragraphs) > 1:
        cell.paragraphs[-1]._element.getparent().remove(cell.paragraphs[-1]._element)

    anchor = template_para._element
    for line in lines[1:]:
        new_p = deepcopy(template_para._element)
        _set_paragraph_text_keep_markup(new_p, line, keep_bold=False)
        anchor.addnext(new_p)
        anchor = new_p


def _normalize_multiline_content_cells(doc: Document) -> None:
    """Multilínea y listas: solo texto, formato de la plantilla."""
    if not doc.tables:
        return
    for row in doc.tables[0].rows:
        if len(row.cells) < 2:
            continue
        label = _cell_label_text(row.cells[0])
        right = row.cells[1]
        if "DETALLE DE PRUEBAS" in label:
            continue

        split_inline = "REFERENCI" in label
        strip_bullets = (
            "REFERENCI" in label
            or "REMEDIACI" in label
            or "AFECTAD" in label
            or "SISTEMA" in label
        )
        multiline = "\n" in (right.text or "")

        if strip_bullets or multiline:
            _expand_cell_lines_from_template(right, split_inline_bullets=split_inline)


def _set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for child in list(tc_pr):
        if child.tag == qn("w:shd"):
            tc_pr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _cell_label_text(cell) -> str:
    raw = cell.text or ""
    return raw.replace("\r", "").replace("\x07", "").replace("\xa0", " ").strip().upper()


def _format_risk_cell(cell, severity: str) -> None:
    """FormatearCeldaNivelRiesgo — etiqueta INAI con fondo y fuente."""
    label = normalize_severity_label(severity).upper()
    fill_hex, font_color = SEVERITY_INAI_STYLE.get(
        label, ("E7E6E6", FONT_BLACK)
    )
    _set_cell_shading(cell, fill_hex)
    for para in cell.paragraphs:
        for run in list(para.runs):
            run.text = ""
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    run = para.add_run(label)
    run.bold = True
    run.font.color.rgb = font_color


def _apply_cyb001_table_formatting(
    doc: Document,
    replacements: dict[str, str],
    severity: str,
    title_cell,
    severity_cell,
) -> None:
    """Solo ajustes que la plantilla no cubre: severidad INAI y etiquetas blancas."""
    if not doc.tables:
        return
    table = doc.tables[0]
    if severity_cell is not None:
        _format_risk_cell(severity_cell, severity)

    detalle_idx = _find_detalle_header_row_idx(table)
    _format_table_label_cells(table, through_row_idx=detalle_idx)

    for row_idx, row in enumerate(table.rows):
        if not row.cells:
            continue
        if detalle_idx is not None and row_idx > detalle_idx:
            break
        left = row.cells[0]
        if _is_detalle_header_cell(left):
            _format_label_cell_white(left)
            continue
        if not _row_is_label_value_pair(row):
            continue
        if _cell_is_label_header(left):
            _format_label_cell_white(left)

    _format_security_detail_content_cells(doc)
    _format_table_body_content_plain(doc)
    # Siempre al final: el paso anterior pone «auto» en celdas de contenido y no debe
    # dejar el encabezado DETALLE en negro (fila fusionada de una sola celda).
    _format_security_detail_header_cells_white(doc)
    _eliminar_ultimas_filas_si_es_salida_prueba_seguridad(doc, replacements)
    # Después de todo: la fila de contenido DETALLE no debe heredar blanco/negrita.
    _format_security_detail_content_row_final(doc)


def _eliminar_ultimas_filas_si_es_salida_prueba_seguridad(
    doc: Document, replacements: dict[str, str]
) -> None:
    """EliminarUltimasFilasSiEsSalidaPruebaSeguridad (macro CYB001)."""
    if not doc.tables:
        return
    salida = (
        replacements.get("«Salidas de herramienta»")
        or replacements.get("«Salida de herramienta»")
        or ""
    ).strip()
    metodo = (replacements.get(METODO_PLACEHOLDER) or "").strip()
    table = doc.tables[0]

    if not salida and not metodo:
        try:
            if len(table.rows) > 0:
                table._tbl.remove(table.rows[-1]._tr)
            if len(table.rows) > 0:
                table._tbl.remove(table.rows[-1]._tr)
        except (IndexError, AttributeError):
            pass
        return

    if not salida:
        for row in table.rows:
            for cell in row.cells:
                if not cell.tables:
                    continue
                nested_tbl = cell.tables[0]._tbl
                nested_tbl.getparent().remove(nested_tbl)
                return


def _field_from_group_or_finding(
    finding: Finding,
    attr: str,
    group_members: Optional[list[Finding]] = None,
) -> str:
    if group_members:
        value = first_non_empty_field(group_members, attr)
        if value is not None:
            return value
    return getattr(finding, attr, None) or ""


def _operational_catalog_field(
    operational_catalog: Optional[dict[str, Any]],
    *cat_keys: str,
    finding: Finding,
    attr: str,
    group_members: Optional[list[Finding]] = None,
) -> str:
    value = catalog_text(operational_catalog, *cat_keys)
    if value:
        return value
    return _field_from_group_or_finding(finding, attr, group_members)


def build_finding_replacements(
    finding: Finding,
    asset: Optional[Asset] = None,
    engagement: Optional[Engagement] = None,
    catalog: Optional[VulnerabilityCatalog] = None,
    extra: Optional[dict[str, str]] = None,
    *,
    merged_componente: Optional[str] = None,
    merged_raw_output: Optional[str] = None,
    group_members: Optional[list[Finding]] = None,
    operational_catalog: Optional[dict[str, Any]] = None,
) -> dict[str, str]:
    """Mapea un hallazgo a marcadores «Columna» estándar."""
    remediation = preprocess_report_field(
        _operational_catalog_field(
            operational_catalog,
            "EspPropuestaRemediacionUnificadaEnRedPrivada",
            "EspPropuestaRemediacionUnificada",
            finding=finding,
            attr="propuesta_remediacion",
            group_members=group_members,
        ),
        strip_bullets=True,
    )
    if not remediation and catalog and catalog.Solution:
        remediation = preprocess_report_field(catalog.Solution, strip_bullets=True)
    elif not remediation and catalog and catalog.EspRemediacion:
        remediation = preprocess_report_field(catalog.EspRemediacion, strip_bullets=True)
    elif not remediation and _field_from_group_or_finding(finding, "descripcion", group_members):
        remediation = preprocess_report_field(
            _field_from_group_or_finding(finding, "descripcion", group_members),
            strip_bullets=True,
        )

    descripcion_raw = _operational_catalog_field(
        operational_catalog,
        "EspDescripcionUnificada",
        finding=finding,
        attr="descripcion",
        group_members=group_members,
    )
    explicacion_source = _operational_catalog_field(
        operational_catalog,
        "EspExplicacionTecnica",
        finding=finding,
        attr="explicacion_tecnica",
        group_members=group_members,
    )
    if not (explicacion_source or "").strip():
        explicacion_source = descripcion_raw or ""
    if _is_markdown(explicacion_source):
        explicacion = _preserve_markdown_field(explicacion_source)
    else:
        explicacion = preprocess_report_field(explicacion_source)
    amenaza = preprocess_report_field(
        _operational_catalog_field(
            operational_catalog,
            "EspAmenazaUnificadaGeneral",
            finding=finding,
            attr="amenaza_ampliada",
            group_members=group_members,
        )
        or (catalog.EspAmenaza if catalog else "")
    )
    descripcion = preprocess_report_field(
        descripcion_raw or (catalog.EspDescripcion if catalog else "")
    )
    if merged_raw_output is not None:
        # Ya incluye prefijos por componente y CYB009 por hallazgo — no re-procesar.
        raw_output = merged_raw_output.strip()
        salidas_markdown = _is_markdown(raw_output)
        tipo_salidas = "markdown" if salidas_markdown and raw_output else "texto plano"
    else:
        raw_output_source = finding.raw_tool_output or ""
        salidas_markdown = _is_markdown(raw_output_source)
        cleaned_salidas = cyb009_salidas_herramienta(
            raw_output_source, for_markdown=salidas_markdown
        )
        if salidas_markdown and cleaned_salidas:
            raw_output = _preserve_markdown_field(cleaned_salidas)
        else:
            raw_output = cleaned_salidas
        tipo_salidas = "markdown" if salidas_markdown and cleaned_salidas else "texto plano"
    metodo_source = _operational_catalog_field(
        operational_catalog,
        "EspMetodoDeteccion",
        finding=finding,
        attr="metodo_deteccion",
        group_members=group_members,
    )
    if _is_markdown(metodo_source):
        metodo_deteccion = _preserve_markdown_field(metodo_source)
    else:
        metodo_deteccion = preprocess_report_field(metodo_source)
    tipo_metodo = "markdown" if _is_markdown(metodo_source) else "texto plano"
    cve_val = _field_from_group_or_finding(finding, "cve", group_members)
    cwe_val = _field_from_group_or_finding(finding, "cwe", group_members)
    refs_parts = []
    if cve_val:
        refs_parts.append(f"CVE: {cve_val}")
    if cwe_val:
        refs_parts.append(f"CWE: {cwe_val}")
    referencias_raw = _operational_catalog_field(
        operational_catalog,
        "References",
        finding=finding,
        attr="referencias",
        group_members=group_members,
    )
    if referencias_raw:
        referencias = preprocess_report_field(
            referencias_raw, strip_bullets=True, split_inline_bullets=True
        )
    else:
        if catalog and catalog.References:
            refs_parts.append(catalog.References)
        referencias = preprocess_report_field(
            "\n".join(refs_parts), strip_bullets=True, split_inline_bullets=True
        )

    tipo_texto = "markdown" if _is_markdown(explicacion_source) else "texto plano"
    sev = finding.severidad.value if finding.severidad else "Medium"
    sev_es = SEVERITY_ES.get(sev, sev)
    sev_inai = normalize_severity_label(sev)
    titulo = preprocess_report_field(
        _operational_catalog_field(
            operational_catalog,
            "EspNombreVulnerabilidadUnificado",
            "StandardVulnerabilityName",
            finding=finding,
            attr="titulo",
            group_members=group_members,
        )
    )

    mapping = {
        "«Título»": titulo,
        "«Nombre de la vulnerabilidad»": titulo,
        "«Nombre de vulnerabilidad»": titulo,
        "«Vulnerabilidad»": titulo,
        "«Descripción»": descripcion,
        "«DESCRIPCIÓN»": descripcion,
        "«Explicación técnica»": explicacion,
        "«Tipo de texto de explicación técnica»": tipo_texto,
        "«Amenaza»": amenaza,
        "«AMENAZA»": amenaza,
        "«Amenaza ampliada»": amenaza,
        "«Propuesta de remediación»": remediation or (finding.propuesta_remediacion or ""),
        "«PROPUESTA DE REMEDIACIÓN»": remediation or (finding.propuesta_remediacion or ""),
        "«Remediación»": remediation,
        "«Solución»": remediation,
        "«Severidad»": sev_inai,
        "«Nivel de riesgo»": sev_inai,
        "«NIVEL DE RIESGO»": sev_inai,
        "«Severidad (español)»": sev_es,
        "«CVE»": cve_val,
        "«CWE»": cwe_val,
        "«CVSS»": str(finding.cvss_score) if finding.cvss_score else "",
        "«CVSS Score»": str(finding.cvss_score) if finding.cvss_score else "",
        "«Vector CVSS»": finding.cvss_vector or "",
        "«Salida de herramienta»": raw_output,
        "«Salidas de herramienta»": raw_output,
        "«Tipo de texto de salidas»": tipo_salidas,
        METODO_PLACEHOLDER: metodo_deteccion,
        "«Tipo de texto de método de detección»": tipo_metodo,
        "«Componente afectado»": (
            merged_componente
            if merged_componente is not None
            else preprocess_report_field(
                finding.componente_afectado or (asset.nombre if asset else ""),
                strip_bullets=True,
            )
        ),
        "«SISTEMA(S) O RUTA(S) AFECTADOS»": (
            merged_componente
            if merged_componente is not None
            else preprocess_report_field(
                finding.componente_afectado or (asset.fqdn or asset.nombre if asset else ""),
                strip_bullets=True,
            )
        ),
        "«Referencias»": referencias,
        "«REFERENCIAS»": referencias,
        "«Estado»": finding.status.value if finding.status else "",
        "«OWASP»": finding.owasp_category or "",
        "«MITRE»": finding.mitre_technique_id or "",
        "«Activo»": asset.nombre if asset else "",
        "«IP»": (asset.ip_publica or asset.ip_privada or "") if asset else "",
        "«FQDN»": asset.fqdn if asset else "",
        "«Cliente»": engagement.cliente if engagement else "",
        "«Fecha»": str(finding.created_at.date()) if finding.created_at else "",
        "«Id»": str(finding.id),
    }

    if catalog:
        mapping.update({
            "«Catálogo»": catalog.DefaultVulnerabilityName or "",
            "«BANOBRAS Categoría»": catalog.BANOBRASCategoryName or "",
            "«BANOBRAS Tipo»": catalog.BANOBRASTipoVulnerabilidad or "",
        })

    if extra:
        for k, v in extra.items():
            key = k if k.startswith("«") else f"«{k}»"
            mapping[key] = v

    return mapping


def generate_single_finding_docx(
    template_path: str,
    output_path: str,
    replacements: dict[str, str],
    severity: str = "Medium",
) -> None:
    shutil.copy2(template_path, output_path)
    doc = Document(output_path)

    explicacion_key = EXPLICACION_PLACEHOLDER
    tipo_explicacion = replacements.get("«Tipo de texto de explicación técnica»", "texto plano")
    tipo_salidas = replacements.get("«Tipo de texto de salidas»", "texto plano")
    tipo_metodo = replacements.get("«Tipo de texto de método de detección»", "texto plano")
    explicacion_val = replacements.get(explicacion_key, "")
    salidas_val = replacements.get("«Salidas de herramienta»", "")
    metodo_val = replacements.get(METODO_PLACEHOLDER, "")

    rich_placeholder_keys: set[str] = set()
    if tipo_explicacion.lower() == "markdown" and explicacion_val:
        rich_placeholder_keys.add(explicacion_key)
    if tipo_salidas.lower() == "markdown" and salidas_val:
        rich_placeholder_keys.update(SALIDAS_PLACEHOLDER_KEYS)
    if tipo_metodo.lower() == "markdown" and metodo_val:
        rich_placeholder_keys.add(METODO_PLACEHOLDER)

    skip_keys = rich_placeholder_keys | set(DETAIL_CONTENT_PLACEHOLDER_KEYS)
    basic_replacements = {
        k: v
        for k, v in replacements.items()
        if k not in skip_keys and not k.startswith(METADATA_PLACEHOLDER_PREFIX)
    }
    # Contenido de celdas de valor: sin heredar negrita del marcador «…» en plantilla.
    _replace_placeholders_in_doc(doc, basic_replacements, keep_bold=False)
    _normalize_multiline_content_cells(doc)

    for detail_key in DETAIL_CONTENT_PLACEHOLDER_KEYS:
        if detail_key in rich_placeholder_keys:
            continue
        detail_val = replacements.get(detail_key, "")
        for para in _iter_all_paragraphs(doc):
            if detail_key in para.text:
                _replace_in_paragraph(
                    para, {detail_key: detail_val}, keep_bold=False
                )

    if explicacion_key in rich_placeholder_keys:
        _insert_rich_markdown_explicacion(doc, explicacion_val)

    if tipo_salidas.lower() == "markdown" and salidas_val:
        _insert_rich_markdown_salidas(doc, salidas_val)

    if METODO_PLACEHOLDER in rich_placeholder_keys:
        _insert_rich_markdown_metodo(doc, metodo_val)

    titulo = (
        replacements.get("«Nombre de vulnerabilidad»")
        or replacements.get("«Nombre de la vulnerabilidad»")
        or replacements.get("«Título»")
        or ""
    )
    title_cell, severity_cell = _get_header_cells(doc)

    _replace_cell_text_preserve_format(title_cell, titulo)
    _apply_cyb001_table_formatting(
        doc, replacements, severity, title_cell, severity_cell
    )
    doc.save(output_path)


def _paragraph_element_text(p_elem) -> str:
    parts: list[str] = []
    for node in p_elem.findall(".//" + qn("w:t")):
        if node.text:
            parts.append(node.text)
    return "".join(parts)


def _paragraph_element_has_page_break(p_elem) -> bool:
    for br in p_elem.findall(".//" + qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def _paragraph_has_page_break(paragraph: Paragraph) -> bool:
    return _paragraph_element_has_page_break(paragraph._element)


def _trim_trailing_empty_body_paragraphs(body) -> None:
    """Quita párrafos vacíos al final del cuerpo (evitan páginas en blanco)."""
    while True:
        children = [c for c in body if c.tag != qn("w:sectPr")]
        if not children:
            break
        last = children[-1]
        if last.tag != qn("w:p"):
            break
        if _paragraph_element_text(last).strip():
            break
        if _paragraph_element_has_page_break(last):
            break
        body.remove(last)


def _insert_before_sect_pr(body, element) -> None:
    """Inserta en el cuerpo antes de w:sectPr (Word exige sectPr al final)."""
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is not None:
        body.insert(list(body).index(sect_pr), element)
    else:
        body.append(element)


def _append_page_break_to_body(body) -> None:
    """Un solo salto de página entre hallazgos (sin párrafo extra de python-docx)."""
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    _insert_before_sect_pr(body, p)


def _collapse_consecutive_page_breaks(doc: Document) -> None:
    """Elimina saltos de página duplicados consecutivos."""
    body = doc.element.body
    prev_break_only = False
    for child in [c for c in body if c.tag != qn("w:sectPr")]:
        if child.tag != qn("w:p"):
            prev_break_only = False
            continue
        text = _paragraph_element_text(child).strip()
        has_break = _paragraph_element_has_page_break(child)
        break_only = has_break and not text
        if break_only and prev_break_only:
            body.remove(child)
            continue
        prev_break_only = break_only


def _strip_trailing_page_breaks(doc: Document) -> None:
    """Elimina saltos de página finales que dejarían una página vacía."""
    body = doc.element.body
    while True:
        children = [child for child in body if child.tag != qn("w:sectPr")]
        if not children:
            break
        last = children[-1]
        if last.tag != qn("w:p"):
            break
        text = _paragraph_element_text(last).strip()
        has_break = _paragraph_element_has_page_break(last)
        if text or not has_break:
            break
        body.remove(last)


MERGE_CHUNK_SIZE = 25


def merge_docx_files(input_paths: list[str], output_path: str) -> None:
    """Fusiona documentos con un salto de página entre hallazgos."""
    if not input_paths:
        raise ValueError("No hay documentos para fusionar")
    master = Document(input_paths[0])
    body = master.element.body
    _trim_trailing_empty_body_paragraphs(body)

    for path in input_paths[1:]:
        _append_page_break_to_body(body)
        sub = Document(path)
        elems = [deepcopy(e) for e in sub.element.body if e.tag != qn("w:sectPr")]
        while elems and elems[0].tag == qn("w:p") and not _paragraph_element_text(elems[0]).strip():
            elems.pop(0)
        for element in elems:
            _insert_before_sect_pr(body, element)
        _trim_trailing_empty_body_paragraphs(body)

    _collapse_consecutive_page_breaks(master)
    _strip_trailing_page_breaks(master)
    master.save(output_path)


def merge_docx_files_chunked(
    input_paths: list[str],
    output_path: str,
    chunk_size: int = MERGE_CHUNK_SIZE,
) -> None:
    """Fusiona lotes grandes en trozos para reducir memoria y tiempo."""
    if not input_paths:
        raise ValueError("No hay documentos para fusionar")
    if len(input_paths) <= chunk_size:
        merge_docx_files(input_paths, output_path)
        return

    with tempfile.TemporaryDirectory() as tmp:
        chunk_paths: list[str] = []
        for i in range(0, len(input_paths), chunk_size):
            chunk = input_paths[i : i + chunk_size]
            chunk_out = str(Path(tmp) / f"chunk_{i // chunk_size}.docx")
            merge_docx_files(chunk, chunk_out)
            chunk_paths.append(chunk_out)
        merge_docx_files(chunk_paths, output_path)


class DocxReportGenerator:
    def __init__(self, storage_dir: str = "storage/reports"):
        self.storage_dir = Path(storage_dir)
        self.storage_dir.mkdir(parents=True, exist_ok=True)

    def generate_batch(
        self,
        template_path: str,
        findings: list[Finding],
        assets_map: dict[uuid.UUID, Asset],
        engagement: Optional[Engagement] = None,
        catalogs_map: Optional[dict[int, VulnerabilityCatalog]] = None,
        job_id: Optional[uuid.UUID] = None,
        db: Optional[Any] = None,
    ) -> dict[str, Any]:
        job_id = job_id or uuid.uuid4()
        work_dir = self.storage_dir / str(job_id)
        work_dir.mkdir(parents=True, exist_ok=True)

        individual_paths: list[str] = []
        catalogs_map = catalogs_map or {}
        grouped_rows = prepare_grouped_rows_for_report(findings, assets_map=assets_map)
        members_by_id = {f.id: f for f in findings}

        from app.services.vulns_catalog_lookup import resolve_operational_catalog_for_finding

        with tempfile.TemporaryDirectory() as tmp:
            for i, row in enumerate(grouped_rows, start=1):
                finding = row.representative
                group_members = [members_by_id[mid] for mid in row.member_ids]
                asset = assets_map.get(finding.asset_id) if finding.asset_id else None
                catalog = catalogs_map.get(finding.catalog_id) if finding.catalog_id else None
                operational_catalog = (
                    resolve_operational_catalog_for_finding(
                        db, finding, group_members=group_members
                    )
                    if db is not None
                    else None
                )
                replacements = build_finding_replacements(
                    finding,
                    asset=asset,
                    engagement=engagement,
                    catalog=catalog,
                    merged_componente=row.merged_componente_afectado,
                    merged_raw_output=row.merged_raw_tool_output,
                    group_members=group_members,
                    operational_catalog=operational_catalog,
                )
                sev = finding.severidad.value if finding.severidad else "Medium"
                tmp_path = str(Path(tmp) / f"doc_{i}.docx")
                generate_single_finding_docx(template_path, tmp_path, replacements, severity=sev)
                final_path = str(work_dir / f"Tabla_{i}.docx")
                shutil.copy2(tmp_path, final_path)
                individual_paths.append(final_path)

            consolidated_name = "Tablas_detalles_vulnerabilidades.docx"
            consolidated_path = str(work_dir / consolidated_name)
            merge_docx_files_chunked(individual_paths, consolidated_path)

        return {
            "job_id": str(job_id),
            "consolidated_path": consolidated_path,
            "individual_paths": individual_paths,
            "findings_count": len(findings),
            "grouped_count": len(grouped_rows),
        }


docx_report_generator = DocxReportGenerator()
